#!/usr/bin/env python3
"""
Generate a Chinese Word (.docx) report from the apibench CSV output.

Usage:
    python3 gen_docx.py <path_to_zh_csv>

Output:
    Same directory as the CSV, with .csv replaced by .docx

Requires: pip install python-docx
"""

import csv
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

CATEGORY_ORDER = [
    "健康检查", "资源管理", "目录树", "文件下载", "文件预览", "文件上传",
    "粘贴/任务", "搜索", "分享", "用户管理", "仓库管理", "权限管理",
    "MD5校验", "外部存储", "回调", "媒体",
]


def set_cell_font(cell, text, size=8, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.bold = bold


def shade_cell(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color,
    })
    shading.append(shd)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def fmt_ms(val):
    if not val:
        return "-"
    v = float(val)
    if v < 1:
        return f"{v*1000:.0f}\u00b5s"
    elif v < 1000:
        return f"{v:.2f}ms"
    else:
        return f"{v/1000:.2f}s"


def suggest_timeout_str(p95_ms, stream):
    if not p95_ms:
        return "1s"
    v = float(p95_ms)
    mult = 5.0 if stream else 3.0
    suggested = v * mult / 1000.0
    if suggested < 1.0:
        suggested = 1.0
    suggested = round(suggested)
    if suggested < 1:
        suggested = 1
    return f"{suggested}s"


def add_styled_paragraph(doc, label, value):
    p = doc.add_paragraph()
    run_l = p.add_run(f"{label}：")
    run_l.bold = True
    run_l.font.size = Pt(10)
    run_l.font.name = "Microsoft YaHei"
    run_l._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run_v = p.add_run(value)
    run_v.font.size = Pt(10)
    run_v.font.name = "Microsoft YaHei"
    run_v._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 gen_docx.py <path_to_zh_csv>", file=sys.stderr)
        sys.exit(1)

    csv_path = sys.argv[1]
    docx_path = csv_path.replace(".csv", ".docx")
    if docx_path == csv_path:
        docx_path = csv_path + ".docx"

    with open(csv_path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    title = doc.add_heading("Files 服务 API 响应时间基准测试报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # --- Summary ---
    add_heading_styled(doc, "一、测试概要", level=1)

    total = len(rows)
    benchmarked = sum(1 for r in rows
                      if r.get("跳过") == "否" and r.get("阶段") == "0"
                      and r.get("备注") not in ("清理",))
    setup_count = sum(1 for r in rows if r.get("备注") == "前置")
    cleanup_count = sum(1 for r in rows if r.get("备注") == "清理")
    skip_count = sum(1 for r in rows if r.get("跳过") == "是")
    error_count = sum(1 for r in rows if r.get("错误"))

    summary_data = [
        ("指标", "数量"),
        ("接口总数", str(total)),
        ("已测试", str(benchmarked)),
        ("前置准备", str(setup_count)),
        ("清理操作", str(cleanup_count)),
        ("跳过（不安全）", str(skip_count)),
        ("错误", str(error_count)),
    ]
    summary_table = doc.add_table(rows=len(summary_data), cols=2, style="Light Grid Accent 1")
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(summary_data):
        bold = i == 0
        set_cell_font(summary_table.rows[i].cells[0], k, size=10, bold=bold)
        set_cell_font(summary_table.rows[i].cells[1], v, size=10, bold=bold)
        if i == 0:
            for c in summary_table.rows[i].cells:
                shade_cell(c, "1F4E79")
                c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # --- Per-category tables ---
    add_heading_styled(doc, "二、各分类接口响应时间", level=1)

    seen_cats = []
    for r in rows:
        cat = r.get("分类", "")
        if cat and cat not in seen_cats:
            seen_cats.append(cat)

    ordered_cats = [c for c in CATEGORY_ORDER if c in seen_cats]
    for c in seen_cats:
        if c not in ordered_cats:
            ordered_cats.append(c)

    for cat in ordered_cats:
        cat_rows = [r for r in rows if r.get("分类") == cat]
        if not cat_rows:
            continue

        add_heading_styled(doc, cat, level=2)

        headers = ["方法", "路径", "说明", "请求大小", "平均", "P50", "P95",
                    "首字节", "最小", "最大", "状态码", "备注"]
        table = doc.add_table(rows=1 + len(cat_rows), cols=len(headers),
                              style="Light Grid Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            set_cell_font(cell, h, size=8, bold=True)
            shade_cell(cell, "1F4E79")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for i, r in enumerate(cat_rows, start=1):
            desc = r.get("说明", "")
            if len(desc) > 35:
                desc = desc[:32] + "..."

            req_bytes = r.get("请求大小(字节)", "0")
            try:
                rb = int(req_bytes)
                if rb == 0:
                    req_size = "-"
                elif rb < 1024:
                    req_size = f"{rb}B"
                elif rb < 1024 * 1024:
                    req_size = f"{rb/1024:.1f}KB"
                else:
                    req_size = f"{rb/(1024*1024):.1f}MB"
            except ValueError:
                req_size = "-"

            vals = [
                r.get("方法", ""),
                (r.get("路径", ""))[:40],
                desc,
                req_size,
                fmt_ms(r.get("平均(ms)")),
                fmt_ms(r.get("P50(ms)")),
                fmt_ms(r.get("P95(ms)")),
                fmt_ms(r.get("首字节均值(ms)")),
                fmt_ms(r.get("最小(ms)")),
                fmt_ms(r.get("最大(ms)")),
                r.get("状态码", "-") or "-",
                r.get("备注", ""),
            ]
            for j, v in enumerate(vals):
                align = (WD_ALIGN_PARAGRAPH.LEFT
                         if j in (1, 2, 11) else WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_font(table.rows[i].cells[j], v, size=7, align=align)

    # --- Timeout recommendations ---
    add_heading_styled(doc, "三、超时建议", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "基于 P95 延迟乘以安全系数（普通接口 3 倍，流式接口 5 倍）计算，"
        "所有建议超时值均设定为不低于 1 秒。"
    )
    run.font.size = Pt(9)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    bench_rows = [r for r in rows
                  if r.get("跳过") == "否"
                  and r.get("平均(ms)")
                  and r.get("备注") not in ("前置", "清理", "依赖跳过")]
    if bench_rows:
        headers = ["分类", "方法", "路径", "P95", "建议超时"]
        t = doc.add_table(rows=1 + len(bench_rows), cols=len(headers),
                          style="Light Grid Accent 1")
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(headers):
            cell = t.rows[0].cells[j]
            set_cell_font(cell, h, size=8, bold=True)
            shade_cell(cell, "1F4E79")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, r in enumerate(bench_rows, start=1):
            p95 = fmt_ms(r.get("P95(ms)"))
            is_stream = r.get("流式") == "是"
            timeout = suggest_timeout_str(r.get("P95(ms)"), is_stream)
            vals = [r.get("分类", ""), r.get("方法", ""),
                    (r.get("路径", ""))[:40], p95, timeout]
            for j, v in enumerate(vals):
                align = (WD_ALIGN_PARAGRAPH.LEFT
                         if j == 2 else WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_font(t.rows[i].cells[j], v, size=7, align=align)

    # --- Skipped route analysis ---
    skip_rows = [r for r in rows if r.get("跳过") == "是"
                 and r.get("分析建议")]
    if skip_rows:
        add_heading_styled(doc, "四、跳过的接口分析", level=1)
        for r in skip_rows:
            add_heading_styled(doc, f"{r['方法']} {r['路径']}", level=3)
            p = doc.add_paragraph()
            run = p.add_run(f"跳过原因：{r.get('跳过原因', '未知')}")
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

            rec = r.get("分析建议", "")
            if rec:
                p2 = doc.add_paragraph()
                run2 = p2.add_run(f"分析建议：{rec}")
                run2.font.size = Pt(9)
                run2.font.name = "Microsoft YaHei"
                run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.save(docx_path)
    print(f"Word 报告已生成: {docx_path}")


if __name__ == "__main__":
    main()
